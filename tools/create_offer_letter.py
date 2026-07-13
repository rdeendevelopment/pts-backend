from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE_IMAGE = Path('/private/tmp/rdeens_offer_inspect/word/media/image1.png')
ASSET_DIR = ROOT / 'generated-offer-assets'
OUTPUT = ROOT / 'Rdeens_Offer_Letter_Ismail_Daniyal_Professional.docx'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def add_bottom_border(cell, color='D9E2F3', size='8'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:color'), color)
    borders.append(bottom)


def add_run(paragraph, text, bold=False, size=9.4, color='222222', italic=False, font='Arial'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_body_paragraph(doc, parts, after=4, before=0, line=1.05):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for text, bold in parts:
        add_run(p, text, bold=bold)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, size=9.7, color='203864')
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, item, size=9.2)


def prepare_brand_assets():
    ASSET_DIR.mkdir(exist_ok=True)
    page = Image.open(SOURCE_IMAGE)
    logo = page.crop((105, 75, 790, 345))

    # Recreate the reference document's geometric header using Rdeens colors.
    header = Image.new('RGB', (1600, 340), 'white')
    draw = ImageDraw.Draw(header)
    draw.rectangle((0, 0, 1600, 90), fill='#293E99')
    draw.polygon([(900, 90), (1600, 90), (1600, 180), (1060, 180)], fill='#4FC3D4')
    header.paste(logo.resize((610, 240)), (40, 82))
    header.save(ASSET_DIR / 'header.png', quality=95)

    # Match the reference footer: contact strip above a strong color band.
    footer = Image.new('RGB', (1600, 165), 'white')
    fdraw = ImageDraw.Draw(footer)
    fdraw.polygon([(0, 78), (355, 78), (520, 165), (0, 165)], fill='#4FC3D4')
    fdraw.rectangle((500, 128, 1600, 165), fill='#293E99')
    font_path = '/System/Library/Fonts/Supplemental/Arial.ttf'
    font = ImageFont.truetype(font_path, 26)
    small = ImageFont.truetype(font_path, 19)
    fdraw.text((470, 67), 'info@rdeens.com', font=font, fill='#444444')
    fdraw.text((820, 67), '+92 300 0524756', font=font, fill='#444444')
    fdraw.text((1160, 67), 'www.rdeens.com', font=font, fill='#444444')
    fdraw.text((1010, 135), 'Din Garden, Chiniot.', font=small, fill='white')
    footer.save(ASSET_DIR / 'footer.png', quality=95)


def build_document():
    prepare_brand_assets()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.68)
    section.bottom_margin = Inches(0.86)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10.6)
    normal.paragraph_format.space_after = Pt(5)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    hp.add_run().add_picture(str(ASSET_DIR / 'header.png'), width=Inches(8.27))

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.add_run().add_picture(str(ASSET_DIR / 'footer.png'), width=Inches(8.27))

    add_body_paragraph(doc, [('Dear Ismail,', False)], before=4, after=15, line=1.15)
    add_body_paragraph(doc, [
        ('Following our discussion, we are pleased to offer you the position of ', False),
        ('AI / RAG Chatbot Developer', True), (' at ', False), ('Rdeens', True),
        (', commencing on ', False), ('15 July 2026', True), ('.', False)
    ], after=9, line=1.18)
    add_body_paragraph(doc, [
        ('We look forward to having you join our team and contribute to the development of reliable AI systems, chatbots, and automation solutions while continuing to grow your professional skills.', False)
    ], after=9, line=1.18)

    add_heading(doc, 'Employment Details:')
    add_bullets(doc, [
        'Role: AI / RAG Chatbot Developer',
        'Joining Date: 15 July 2026',
        'Monthly Compensation: PKR 40,000',
        'Working Commitment: Four hours per day',
        'Work Arrangement: Hybrid'
    ])

    add_heading(doc, 'Key Responsibilities:')
    add_bullets(doc, [
        'Develop and maintain RAG-based AI systems and AI chatbot solutions.',
        'Design and implement AI automation workflows and related technical features.',
        'Research, implement, test, document, and improve AI-powered functionality.',
        'Complete other relevant technical assignments reasonably required by Rdeens.'
    ])

    add_heading(doc, 'Terms and Conditions:')
    terms = [
        ('1. Daily Reporting: ', 'Push completed work to the assigned GitHub repository, submit a daily report covering completed work, pending tasks, and blockers, and communicate technical issues promptly.'),
        ('2. Work Arrangement: ', 'This is a hybrid role. You may work remotely; however, you will be expected to attend the office whenever reasonably required by the company.'),
        ('3. Confidentiality: ', 'All project data, source code, credentials, client information, documents, business information, and internal discussions must remain strictly confidential and may not be shared without prior written authorization from Rdeens.'),
        ('4. Ownership of Work: ', 'Unless otherwise agreed in writing, all code, AI systems, chatbots, automation workflows, documentation, and technical assets created during your engagement will remain the exclusive property of Rdeens.'),
        ('5. Performance Review: ', 'Your performance, commitment, technical capability, communication, and overall contribution will be reviewed during the probationary period.')
    ]
    for label, detail in terms:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.02
        add_run(p, label, bold=True, size=9.2)
        add_run(p, detail, size=9.2)

    add_body_paragraph(doc, [
        ('Please sign and return a copy of this letter to confirm your acceptance. We look forward to welcoming you to Rdeens and building a productive, long-term professional relationship.', False)
    ], before=5, after=9, line=1.12)

    sig = doc.add_table(rows=1, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig.autofit = False
    sig.columns[0].width = Inches(3.2)
    sig.columns[1].width = Inches(3.2)
    for cell in sig.rows[0].cells:
        set_cell_margins(cell, top=20, start=60, bottom=20, end=60)
    left = sig.cell(0, 0)
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    add_run(lp, 'Warm regards,', size=10.4)
    np = left.add_paragraph()
    np.paragraph_format.space_after = Pt(0)
    add_run(np, 'Usama', size=10.4)
    tp = left.add_paragraph()
    tp.paragraph_format.space_after = Pt(0)
    add_run(tp, 'Authorized Signatory, Rdeens', size=10.4)

    right = sig.cell(0, 1)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.paragraph_format.space_after = Pt(0)
    add_run(rp, 'Usama', size=20, color='203864', italic=True, font='Snell Roundhand')
    rline = right.add_paragraph()
    rline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rline.paragraph_format.space_after = Pt(0)
    add_run(rline, '________________________', size=9.0, color='666666')
    rlabel = right.add_paragraph()
    rlabel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rlabel.paragraph_format.space_after = Pt(0)
    add_run(rlabel, 'Authorized Signature', size=8.8, color='666666')

    doc.core_properties.title = 'Offer of Employment — Ismail Daniyal'
    doc.core_properties.subject = 'AI / RAG Chatbot Developer Offer Letter'
    doc.core_properties.author = 'Rdeens'
    doc.core_properties.keywords = 'Rdeens, offer letter, Ismail Daniyal'
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    build_document()
